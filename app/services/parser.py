from __future__ import annotations

import json
import re
import traceback
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass
from pathlib import Path

from app.services.storage import ensure_storage


@dataclass
class ParsedBlock:
    content: str
    block_type: str = "text"
    heading_path: list[str] | None = None
    page: int | None = None
    bbox: list[float] | None = None
    source_method: str = "unknown"


@dataclass
class ParseQuality:
    status: str
    total_pages: int = 0
    text_pages: int = 0
    empty_pages: int = 0
    total_blocks: int = 0
    avg_blocks_per_page: float = 0
    short_block_ratio: float = 0
    source_method: str = "unknown"
    notes: list[str] | None = None


def _parse_markdown(path: Path, source_method: str = "markdown") -> list[ParsedBlock]:
    text = path.read_text(encoding="utf-8", errors="replace")
    blocks: list[ParsedBlock] = []
    headings: list[str] = []
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            value = "\n".join(buffer).strip()
            if value:
                kind = "table" if "|" in value and "---" in value else "text"
                blocks.append(ParsedBlock(value, kind, headings.copy(), source_method=source_method))
            buffer.clear()

    for line in text.splitlines():
        if line.startswith("#"):
            flush()
            level = len(line) - len(line.lstrip("#"))
            title = line[level:].strip()
            headings[:] = headings[: level - 1] + [title]
        elif not line.strip():
            flush()
        else:
            buffer.append(line)
    flush()
    return blocks


def _parse_docx(path: Path) -> list[ParsedBlock]:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    blocks: list[ParsedBlock] = []
    headings: list[str] = []
    for paragraph in doc.paragraphs:
        value = paragraph.text.strip()
        if not value:
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            try:
                level = int(paragraph.style.name.split()[-1])
            except ValueError:
                level = 1
            headings[:] = headings[: level - 1] + [value]
        else:
            blocks.append(ParsedBlock(value, "text", headings.copy(), source_method="docx"))
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            markdown = "| " + " | ".join(rows[0]) + " |\n"
            markdown += "| " + " | ".join("---" for _ in rows[0]) + " |\n"
            markdown += "\n".join("| " + " | ".join(row) + " |" for row in rows[1:])
            blocks.append(ParsedBlock(markdown, "table", headings.copy(), source_method="docx"))
    return blocks


def _toc_paths(toc: list[list]) -> list[tuple[int, list[str]]]:
    stack: list[str] = []
    paths: list[tuple[int, list[str]]] = []
    for row in toc:
        if len(row) < 3:
            continue
        level, title, page = int(row[0]), str(row[1]).strip(), int(row[2])
        if not title or page < 1:
            continue
        stack = stack[: max(0, level - 1)] + [title]
        paths.append((page, stack.copy()))
    return paths


def _heading_for_page(paths: list[tuple[int, list[str]]], page_number: int) -> list[str]:
    heading: list[str] = []
    for start_page, path in paths:
        if start_page > page_number:
            break
        heading = path
    return heading


def _normalize_noise_candidate(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip().lower()


def _is_page_number(value: str) -> bool:
    text = value.strip()
    return bool(re.fullmatch(r"[-–—]?\s*(?:page\s*)?\d{1,4}\s*[-–—]?", text, flags=re.IGNORECASE))


def _is_short_noise(value: str) -> bool:
    text = re.sub(r"\s+", "", value)
    if not text:
        return True
    if _is_page_number(value):
        return True
    return len(text) <= 2 and not re.search(r"[\u3400-\u9fffA-Za-z0-9]", text)


def _quality_from_stats(
    *,
    total_pages: int,
    text_pages: int,
    total_blocks: int,
    short_blocks: int,
    source_method: str,
) -> ParseQuality:
    empty_pages = max(total_pages - text_pages, 0)
    avg_blocks = round(total_blocks / total_pages, 2) if total_pages else 0
    short_ratio = round(short_blocks / total_blocks, 4) if total_blocks else 0
    text_page_ratio = text_pages / total_pages if total_pages else 0
    notes: list[str] = []
    status = "GOOD"

    if total_pages == 0:
        status = "POOR"
        notes.append("PDF 没有可读取页面。")
    elif total_blocks == 0 or text_page_ratio < 0.4:
        status = "POOR"
        notes.append("大量页面没有可提取文本，疑似扫描 PDF 或文本层缺失。")
    elif text_page_ratio < 0.85 or avg_blocks < 1.5 or short_ratio > 0.35:
        status = "WARNING"
        if text_page_ratio < 0.85:
            notes.append("部分页面没有可提取文本。")
        if avg_blocks < 1.5:
            notes.append("平均每页文本块偏少，可能存在解析不完整。")
        if short_ratio > 0.35:
            notes.append("短文本块比例偏高，可能包含页眉页脚或版式噪声。")

    return ParseQuality(
        status=status,
        total_pages=total_pages,
        text_pages=text_pages,
        empty_pages=empty_pages,
        total_blocks=total_blocks,
        avg_blocks_per_page=avg_blocks,
        short_block_ratio=short_ratio,
        source_method=source_method,
        notes=notes,
    )


def _write_parse_metadata(output_dir: Path, quality: ParseQuality) -> None:
    (output_dir / "metadata.json").write_text(json.dumps(asdict(quality), ensure_ascii=False, indent=2), encoding="utf-8")


def _write_error_book(
    document_id: int,
    filename: str,
    stage: str,
    quality: ParseQuality | None = None,
    error: str | None = None,
) -> None:
    root = ensure_storage()
    payload = {
        "document_id": document_id,
        "filename": filename,
        "stage": stage,
        "quality": asdict(quality) if quality else None,
        "error": error,
    }
    (root / "error-book" / f"document-{document_id}-{stage}.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def _render_pdf_pages(path: Path, document_id: int, output_dir: Path) -> tuple[list[ParsedBlock], ParseQuality]:
    import fitz

    root = ensure_storage()
    page_dir = root / "pages" / str(document_id)
    page_dir.mkdir(parents=True, exist_ok=True)
    raw_blocks_by_page: dict[int, list[tuple[str, list[float], float]]] = defaultdict(list)
    total_pages = 0
    text_pages = 0
    short_blocks = 0
    render_errors = 0
    with fitz.open(path) as pdf:
        total_pages = len(pdf)
        toc = pdf.get_toc(simple=True)
        toc_paths = _toc_paths(toc)
        if toc:
            (output_dir / "toc.json").write_text(json.dumps(toc, ensure_ascii=False, indent=2), encoding="utf-8")
        for page_index, page in enumerate(pdf):
            page_number = page_index + 1
            page_image_path = page_dir / f"{page_number}.png"
            if not page_image_path.exists():
                try:
                    pix = page.get_pixmap(matrix=fitz.Matrix(1.2, 1.2), alpha=False)
                    pix.save(page_image_path)
                    del pix
                except Exception:
                    render_errors += 1
            page_has_text = False
            page_height = float(page.rect.height)
            for raw in page.get_text("blocks", sort=True):
                content = raw[4].strip()
                if content:
                    page_has_text = True
                    if _is_short_noise(content):
                        short_blocks += 1
                    raw_blocks_by_page[page_number].append((content, list(raw[:4]), page_height))
            if page_has_text:
                text_pages += 1

        repeated_edge_texts: Counter[str] = Counter()
        for page_blocks in raw_blocks_by_page.values():
            seen_on_page: set[str] = set()
            for content, bbox, page_height in page_blocks:
                top_or_bottom = bbox[1] < page_height * 0.09 or bbox[3] > page_height * 0.91
                normalized = _normalize_noise_candidate(content)
                if top_or_bottom and normalized and len(normalized) <= 120:
                    seen_on_page.add(normalized)
            repeated_edge_texts.update(seen_on_page)

        repeated_threshold = max(3, int(total_pages * 0.35)) if total_pages else 3
        repeated_noise = {text for text, count in repeated_edge_texts.items() if count >= repeated_threshold}

        blocks: list[ParsedBlock] = []
        for page_number in sorted(raw_blocks_by_page):
            heading_path = _heading_for_page(toc_paths, page_number)
            for content, bbox, page_height in raw_blocks_by_page[page_number]:
                normalized = _normalize_noise_candidate(content)
                top_or_bottom = bbox[1] < page_height * 0.09 or bbox[3] > page_height * 0.91
                if _is_page_number(content):
                    continue
                if top_or_bottom and normalized in repeated_noise:
                    continue
                if _is_short_noise(content):
                    continue
                blocks.append(ParsedBlock(content, "text", heading_path.copy(), page_number, bbox, "pdf_text"))

    quality = _quality_from_stats(
        total_pages=total_pages,
        text_pages=text_pages,
        total_blocks=len(blocks),
        short_blocks=short_blocks,
        source_method="pdf_text",
    )
    if render_errors:
        quality.notes = [*(quality.notes or []), f"{render_errors} 页页面截图生成失败；文本解析已保留。"]
        if quality.status == "GOOD":
            quality.status = "WARNING"
    return blocks, quality


def _docling_markdown(path: Path) -> tuple[str, str]:
    try:
        from docling.document_converter import DocumentConverter
    except ImportError as exc:
        raise RuntimeError("当前 Demo 未安装 Docling/OCR 扩展；请优先使用带文本层的原生 PDF。") from exc

    result = DocumentConverter().convert(path)
    document = result.document
    markdown = document.export_to_markdown()
    try:
        structured = document.export_to_dict()
    except AttributeError:
        structured = {"markdown": markdown}
    return markdown, json.dumps(structured, ensure_ascii=False, indent=2)


def parse_document(path: Path, document_id: int, parse_mode: str = "auto") -> list[ParsedBlock]:
    suffix = path.suffix.lower()
    root = ensure_storage()
    output_dir = root / "parsed" / str(document_id)
    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        if suffix in {".md", ".markdown", ".txt"}:
            blocks = _parse_markdown(path)
            quality = ParseQuality("GOOD", total_blocks=len(blocks), source_method="markdown", notes=[])
            _write_parse_metadata(output_dir, quality)
        elif suffix == ".docx":
            blocks = _parse_docx(path)
            quality = ParseQuality(
                "GOOD" if blocks else "POOR",
                total_blocks=len(blocks),
                source_method="docx",
                notes=[] if blocks else ["DOCX 未提取到文本块。"],
            )
            _write_parse_metadata(output_dir, quality)
            if quality.status == "POOR":
                _write_error_book(document_id, path.name, "parse-quality", quality)
        elif suffix == ".pdf":
            blocks, quality = _render_pdf_pages(path, document_id, output_dir)
            use_docling = parse_mode in {"ocr", "vlm"} or (parse_mode == "auto" and not blocks)
            if use_docling:
                try:
                    markdown, structured = _docling_markdown(path)
                    (output_dir / "document.md").write_text(markdown, encoding="utf-8")
                    (output_dir / "document.json").write_text(structured, encoding="utf-8")
                    temp = output_dir / "document.md"
                    docling_blocks = _parse_markdown(temp, source_method="docling_markdown")
                    if docling_blocks:
                        blocks = docling_blocks
                        quality = ParseQuality(
                            "WARNING",
                            total_pages=quality.total_pages,
                            text_pages=quality.text_pages,
                            empty_pages=quality.empty_pages,
                            total_blocks=len(blocks),
                            avg_blocks_per_page=round(len(blocks) / quality.total_pages, 2) if quality.total_pages else 0,
                            short_block_ratio=quality.short_block_ratio,
                            source_method="docling_markdown",
                            notes=["PyMuPDF 文本提取不足，已使用 Docling Markdown 兜底。"],
                        )
                except Exception as exc:
                    if not blocks:
                        quality.source_method = "ocr_pending"
                        quality.status = "POOR"
                        quality.notes = [*(quality.notes or []), f"Docling/OCR 兜底失败：{type(exc).__name__}: {exc}"]
                        _write_error_book(document_id, path.name, "parse-failed", quality, traceback.format_exc())
            if quality.status == "POOR" and quality.source_method == "pdf_text":
                quality.source_method = "ocr_pending"
                quality.notes = [*(quality.notes or []), "当前阶段暂不对扫描 PDF 执行完整 OCR，请后续用 OCR/VLM 模式重建。"]
            _write_parse_metadata(output_dir, quality)
            if quality.status == "POOR":
                _write_error_book(document_id, path.name, "parse-quality", quality)
        else:
            raise ValueError(f"不支持的文件类型: {suffix}")
    except Exception:
        quality = ParseQuality("POOR", source_method="parse_failed", notes=["解析过程抛出异常。"])
        _write_parse_metadata(output_dir, quality)
        _write_error_book(document_id, path.name, "parse-failed", quality, traceback.format_exc())
        raise

    (output_dir / "blocks.json").write_text(
        json.dumps([asdict(block) for block in blocks], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return blocks
